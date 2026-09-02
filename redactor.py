"""
PII Redactor - A simple tool to detect and replace Personally Identifiable Information (PII)
in Word Documents (.docx) using Microsoft Presidio and Faker.
"""

import json
import re
from docx import Document
from faker import Faker
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer

# ---------------------------------------------------------
# 1. Setup Faker for generating realistic fake data
# ---------------------------------------------------------
fake = Faker()
Faker.seed(123)  # Seed ensures we get the same fakes if we run it twice

# We use a dictionary to remember our replacements.
# If we see "John Doe" 5 times, we want to replace it with the same fake name every time.
replacement_memory = {}

def get_fake_value(entity_type, original_text):
    """Generates a realistic fake value based on the type of PII."""
    original_text = original_text.strip()
    
    # Check if we have already generated a fake value for this exact text
    memory_key = (entity_type, original_text.lower())
    if memory_key in replacement_memory:
        return replacement_memory[memory_key]
        
    # Generate a new fake value based on the entity type
    if entity_type == "PERSON":
        fake_val = fake.name()
    elif entity_type == "EMAIL_ADDRESS":
        fake_val = fake.email()
    elif entity_type == "PHONE_NUMBER":
        # Format as Indian phone number if appropriate
        if "+91" in original_text or len(original_text.replace(" ", "")) >= 10:
            fake_val = fake.numerify("+91 ##### #####")
        else:
            fake_val = fake.phone_number()
    elif entity_type == "US_SSN":
        fake_val = fake.ssn()
    elif entity_type == "CREDIT_CARD":
        # Generate a visually realistic credit card (spaced blocks of 4)
        cc = fake.credit_card_number(card_type="visa16")
        fake_val = f"{cc[:4]} {cc[4:8]} {cc[8:12]} {cc[12:16]}"
    elif entity_type == "IP_ADDRESS":
        fake_val = fake.ipv4()
    elif entity_type == "DATE_TIME":
        fake_val = fake.date_of_birth(minimum_age=18, maximum_age=70).strftime("%d-%m-%Y")
    elif entity_type in ["LOCATION", "NRP"]:  
        # NRP (Nationality, Religious, Political) often catches company names in Presidio
        fake_val = fake.company() if entity_type == "NRP" else fake.city()
    else:
        fake_val = "[REDACTED]"
        
    # Save it in memory for next time
    replacement_memory[memory_key] = fake_val
    return fake_val

# ---------------------------------------------------------
# 2. Setup Microsoft Presidio Analyzer
# ---------------------------------------------------------
analyzer = AnalyzerEngine()

# Add a custom regex recognizer for Indian Phone Numbers (+91 XXXXXXXXXX)
indian_phone_pattern = Pattern(
    name="indian_phone", 
    regex=r"(\+91[\-\s]?)?[6-9]\d{4}[\-\s]?\d{5}", 
    score=0.85
)
analyzer.registry.add_recognizer(PatternRecognizer(supported_entity="PHONE_NUMBER", patterns=[indian_phone_pattern]))

# Add a custom regex recognizer for Dates (DD/MM/YYYY or DD-Month-YYYY)
date_pattern1 = Pattern(name="date_slash", regex=r"\b\d{2}/\d{2}/\d{4}\b", score=0.8)
date_pattern2 = Pattern(name="date_dash", regex=r"\b\d{1,2}-[A-Za-z]+-\d{4}\b", score=0.8)
analyzer.registry.add_recognizer(PatternRecognizer(supported_entity="DATE_TIME", patterns=[date_pattern1, date_pattern2]))

# ---------------------------------------------------------
# 3. Document Processing Logic
# ---------------------------------------------------------
def redact_paragraph(paragraph):
    """Reads a single paragraph, finds PII, and replaces it with fake data."""
    text = paragraph.text
    if not text.strip():
        return [] # Empty paragraph, skip it

    # Ask Presidio to find the PII entities in this text block
    results = analyzer.analyze(
        text=text,
        language="en",
        entities=["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "US_SSN", "CREDIT_CARD", "IP_ADDRESS", "DATE_TIME", "LOCATION", "NRP"],
        score_threshold=0.4
    )
    
    # Sort results from end to start so replacing text doesn't mess up the indices for earlier replacements
    results.sort(key=lambda x: x.start, reverse=True)
    
    detected_pii_log = []
    redacted_text = text
    
    for result in results:
        original_pii = text[result.start:result.end]
        fake_pii = get_fake_value(result.entity_type, original_pii)
        
        # Splice the fake value into the string
        redacted_text = redacted_text[:result.start] + fake_pii + redacted_text[result.end:]
        
        # Keep track of what we found for the evaluation log
        detected_pii_log.append({
            "type": result.entity_type,
            "original": original_pii,
            "replaced_with": fake_pii
        })

    # If we made any changes, update the actual Word document paragraph
    if detected_pii_log:
        # Clear existing text blocks (runs) to avoid duplicate text
        for run in paragraph.runs:
            run.text = ""
        # Put the new redacted text in the first block
        if paragraph.runs:
            paragraph.runs[0].text = redacted_text
        else:
            paragraph.add_run(redacted_text)
            
    return detected_pii_log

def process_document(input_file, output_file, log_file):
    """Opens the Word doc, processes every paragraph and table, and saves the result."""
    print(f"Reading {input_file}...")
    doc = Document(input_file)
    all_logs = []
    
    # 1. Process standard body paragraphs
    for para in doc.paragraphs:
        all_logs.extend(redact_paragraph(para))
        
    # 2. Process paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_logs.extend(redact_paragraph(para))
                    
    # Save the new document
    print(f"Saving redacted document to {output_file}...")
    doc.save(output_file)
    
    # Save the logs to a JSON file so we can analyze them
    with open(log_file, "w") as f:
        json.dump(all_logs, f, indent=4)
    print(f"Saved logs to {log_file}. Total PII instances redacted: {len(all_logs)}")

if __name__ == "__main__":
    process_document(
        input_file="PII_Assignment_Source_Document_KSH_RHP.docx",
        output_file="Redacted_Output.docx",
        log_file="redaction_log.json"
    )
